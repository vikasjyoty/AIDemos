from __future__ import annotations

import csv
from pathlib import Path
from typing import Dict, List

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader


ParsedSection = Dict[str, object]


def _read_text_file(file_path: Path) -> List[ParsedSection]:
    print(f"[INFO] Parsing text-like file: {file_path.name}")
    text = file_path.read_text(encoding="utf-8", errors="ignore")
    return [{"text": text, "content_type": "text", "metadata": {"parser": "plain_text"}}]


def _parse_html(file_path: Path) -> List[ParsedSection]:
    print(f"[INFO] Parsing HTML file: {file_path.name}")
    html = file_path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "lxml")

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    sections: List[ParsedSection] = []

    visible_text = soup.get_text(separator=" ", strip=True)
    if visible_text:
        sections.append(
            {
                "text": visible_text,
                "content_type": "html_text",
                "metadata": {"parser": "beautifulsoup"},
            }
        )

    table_count = 0
    for table_index, table in enumerate(soup.find_all("table"), start=1):
        rows: List[str] = []
        for row in table.find_all("tr"):
            cells = row.find_all(["th", "td"])
            values = [cell.get_text(" ", strip=True) for cell in cells]
            if any(values):
                rows.append(" | ".join(values))

        table_text = "\n".join(rows).strip()
        if table_text:
            table_count += 1
            sections.append(
                {
                    "text": table_text,
                    "content_type": "table",
                    "metadata": {"parser": "html_table", "table_index": table_index},
                }
            )

    print(f"[OK] HTML parsed with {table_count} table(s) extracted.")
    return sections


def _parse_csv(file_path: Path) -> List[ParsedSection]:
    print(f"[INFO] Parsing CSV file: {file_path.name}")
    rows: List[str] = []
    with file_path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        reader = csv.reader(handle)
        for row in reader:
            if row:
                rows.append(" | ".join(cell.strip() for cell in row))

    table_text = "\n".join(rows)
    return [{"text": table_text, "content_type": "table", "metadata": {"parser": "csv"}}]


def _parse_pdf(file_path: Path) -> List[ParsedSection]:
    print(f"[INFO] Parsing PDF file: {file_path.name}")
    sections: List[ParsedSection] = []

    reader = PdfReader(str(file_path))
    text_parts: List[str] = []
    for page in reader.pages:
        text_parts.append(page.extract_text() or "")

    full_text = "\n".join(text_parts).strip()
    if full_text:
        sections.append(
            {
                "text": full_text,
                "content_type": "pdf_text",
                "metadata": {"parser": "pypdf"},
            }
        )

    # Optional table extraction from PDF pages.
    table_count = 0
    try:
        import pdfplumber

        with pdfplumber.open(str(file_path)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables() or []
                for table_index, table in enumerate(tables, start=1):
                    lines: List[str] = []
                    for row in table:
                        cleaned = [str(cell or "").strip() for cell in row]
                        if any(cleaned):
                            lines.append(" | ".join(cleaned))
                    table_text = "\n".join(lines).strip()
                    if table_text:
                        table_count += 1
                        sections.append(
                            {
                                "text": table_text,
                                "content_type": "table",
                                "metadata": {
                                    "parser": "pdfplumber",
                                    "page": page_index,
                                    "table_index": table_index,
                                },
                            }
                        )
    except ImportError:
        print("[WARN] pdfplumber not installed. PDF text extracted, table extraction skipped.")

    print(f"[OK] PDF parsed with {table_count} table(s) extracted.")
    return sections


def _parse_docx(file_path: Path) -> List[ParsedSection]:
    print(f"[INFO] Parsing DOCX file: {file_path.name}")
    document = Document(str(file_path))
    sections: List[ParsedSection] = []

    paragraph_text = "\n".join(p.text.strip() for p in document.paragraphs if p.text.strip()).strip()
    if paragraph_text:
        sections.append(
            {
                "text": paragraph_text,
                "content_type": "docx_text",
                "metadata": {"parser": "python-docx"},
            }
        )

    table_count = 0
    for table_index, table in enumerate(document.tables, start=1):
        lines: List[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
        table_text = "\n".join(lines).strip()
        if table_text:
            table_count += 1
            sections.append(
                {
                    "text": table_text,
                    "content_type": "table",
                    "metadata": {"parser": "python-docx", "table_index": table_index},
                }
            )

    print(f"[OK] DOCX parsed with {table_count} table(s) extracted.")
    return sections


def parse_document(file_path: Path) -> List[ParsedSection]:
    """Parse a document into sections that can be chunked and embedded."""
    suffix = file_path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return _read_text_file(file_path)
    if suffix in {".html", ".htm"}:
        return _parse_html(file_path)
    if suffix == ".csv":
        return _parse_csv(file_path)
    if suffix == ".pdf":
        return _parse_pdf(file_path)
    if suffix == ".docx":
        return _parse_docx(file_path)

    print(f"[WARN] Unsupported file type skipped: {file_path.name}")
    return []
